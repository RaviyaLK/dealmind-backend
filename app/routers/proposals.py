from fastapi import APIRouter, Depends, HTTPException, Query, BackgroundTasks, UploadFile, File, Form
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from uuid import uuid4
from datetime import datetime
import aiofiles
import json
import os
import re
import tempfile
from app.database import get_db
from app.services.auth import get_current_user
from app.models.proposal import Proposal
from app.models.deal import Deal
from app.models.document import Document
from app.schemas.proposal import (
    ProposalCreate,
    ProposalResponse,
    ProposalReview,
    ProposalExportRequest,
)

router = APIRouter(prefix="/api/proposals", tags=["Proposals"])

# ═══════════════════════════════════════════════════
# PROPOSAL TEMPLATES
# ═══════════════════════════════════════════════════
TEMPLATES = {
    "modern": {
        "name": "Modern",
        "description": "Bold & contemporary with vibrant purple accents",
        "primary_hex": "6B21A8",
        "primary_rgb": (0x6B, 0x21, 0xA8),
        "primary_light_hex": "9B59D0",
        "primary_light_rgb": (0x9B, 0x59, 0xD0),
        "accent_hex": "06B6D4",
        "accent_rgb": (0x06, 0xB6, 0xD4),
        "dark_hex": "1A1A2E",
        "dark_rgb": (0x1A, 0x1A, 0x2E),
        "grey_hex": "4B5563",
        "grey_rgb": (0x4B, 0x55, 0x63),
        "light_grey_hex": "9CA3AF",
        "light_grey_rgb": (0x9C, 0xA3, 0xAF),
        "heading_font": "Calibri",
        "body_font": "Calibri",
        "cover_border_sz": 48,
        "rule_thickness": 2,
        "table_header_bg": "6B21A8",
        "table_header_text_rgb": (0xFF, 0xFF, 0xFF),
        "table_row_alt_bg": "F3E8FF",
        "table_border_hex": "D8B4FE",
        "table_border_inside": "E9D5FF",
    },
    "classic": {
        "name": "Classic",
        "description": "Traditional & formal with navy blue and gold accents",
        "primary_hex": "1E3A5F",
        "primary_rgb": (0x1E, 0x3A, 0x5F),
        "primary_light_hex": "2E5090",
        "primary_light_rgb": (0x2E, 0x50, 0x90),
        "accent_hex": "D4A574",
        "accent_rgb": (0xD4, 0xA5, 0x74),
        "dark_hex": "0F172A",
        "dark_rgb": (0x0F, 0x17, 0x2A),
        "grey_hex": "475569",
        "grey_rgb": (0x47, 0x55, 0x69),
        "light_grey_hex": "94A3B8",
        "light_grey_rgb": (0x94, 0xA3, 0xB8),
        "heading_font": "Times New Roman",
        "body_font": "Calibri",
        "cover_border_sz": 36,
        "rule_thickness": 2,
        "table_header_bg": "1E3A5F",
        "table_header_text_rgb": (0xFF, 0xFF, 0xFF),
        "table_row_alt_bg": "F0F4F8",
        "table_border_hex": "CBD5E1",
        "table_border_inside": "E2E8F0",
    },
    "minimal": {
        "name": "Minimal",
        "description": "Ultra-clean monochrome design with a single accent color",
        "primary_hex": "18181B",
        "primary_rgb": (0x18, 0x18, 0x1B),
        "primary_light_hex": "3F3F46",
        "primary_light_rgb": (0x3F, 0x3F, 0x46),
        "accent_hex": "0891B2",
        "accent_rgb": (0x08, 0x91, 0xB2),
        "dark_hex": "09090B",
        "dark_rgb": (0x09, 0x09, 0x0B),
        "grey_hex": "52525B",
        "grey_rgb": (0x52, 0x52, 0x5B),
        "light_grey_hex": "A1A1AA",
        "light_grey_rgb": (0xA1, 0xA1, 0xAA),
        "heading_font": "Calibri",
        "body_font": "Calibri",
        "cover_border_sz": 12,
        "rule_thickness": 1,
        "table_header_bg": "F4F4F5",
        "table_header_text_rgb": (0x18, 0x18, 0x1B),
        "table_row_alt_bg": "FAFAFA",
        "table_border_hex": "E4E4E7",
        "table_border_inside": "F4F4F5",
    },
}


@router.get("/templates")
def list_templates():
    """Return available proposal templates for the frontend selector."""
    return [
        {
            "id": tid,
            "name": t["name"],
            "description": t["description"],
            "primary_hex": t["primary_hex"],
            "accent_hex": t["accent_hex"],
            "heading_font": t["heading_font"],
            "body_font": t["body_font"],
        }
        for tid, t in TEMPLATES.items()
    ]


@router.get("/", response_model=list[ProposalResponse])
def list_proposals(
    deal_id: str | None = Query(None),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user),
):
    """
    List proposals with optional deal_id filter.
    """
    query = db.query(Proposal)

    if deal_id:
        query = query.filter(Proposal.deal_id == deal_id)

    proposals = query.all()

    return [ProposalResponse.model_validate(prop) for prop in proposals]


@router.post("/", response_model=ProposalResponse)
def create_proposal(
    proposal_data: ProposalCreate,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user),
):
    """
    Create proposal placeholder. Triggers agent later.
    """
    proposal_id = str(uuid4())
    proposal = Proposal(
        id=proposal_id,
        deal_id=proposal_data.deal_id,
        title=proposal_data.title or f"Proposal {proposal_id[:8]}",
        content="Draft - pending agent generation",
        status="draft",
        generated_by=current_user.id,
    )

    db.add(proposal)
    db.commit()
    db.refresh(proposal)

    background_tasks.add_task(trigger_proposal_agent, proposal_id, proposal_data.deal_id)

    return ProposalResponse.model_validate(proposal)


def trigger_proposal_agent(proposal_id: str, deal_id: str):
    """Placeholder for agent trigger - to be implemented."""
    pass


@router.get("/{proposal_id}", response_model=ProposalResponse)
def get_proposal(
    proposal_id: str,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user),
):
    """
    Get proposal detail.
    """
    proposal = db.query(Proposal).filter(Proposal.id == proposal_id).first()

    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    return ProposalResponse.model_validate(proposal)


@router.patch("/{proposal_id}/review", response_model=ProposalResponse)
def submit_proposal_review(
    proposal_id: str,
    review_data: ProposalReview,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user),
):
    """
    Submit review (approve/reject with notes).
    """
    proposal = db.query(Proposal).filter(Proposal.id == proposal_id).first()

    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    proposal.status = review_data.status
    proposal.review_notes = review_data.review_notes
    proposal.reviewed_by = current_user.id

    db.commit()
    db.refresh(proposal)

    return ProposalResponse.model_validate(proposal)


# ═══════════════════════════════════════════════════
# PROPOSAL CHAT — Refine proposal via MCP tools
# ═══════════════════════════════════════════════════
from pydantic import BaseModel
from typing import List, Optional as Opt


class ChatMessage(BaseModel):
    role: str  # "user" or "assistant"
    content: str


class ProposalChatRequest(BaseModel):
    message: str
    history: List[ChatMessage] = []


def _clean_json_block(text: str) -> str:
    """Strip markdown code-block wrappers from a JSON string."""
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*\n?", "", cleaned)
        cleaned = re.sub(r"\n?```\s*$", "", cleaned)
        cleaned = cleaned.strip()
    return cleaned


def _extract_tool_calls(response: str) -> tuple:
    """Try to extract tool_calls from LLM response.

    Returns (tool_calls_list, remaining_text).
    tool_calls_list is [] if no tool calls found.
    """
    cleaned = _clean_json_block(response)

    # Try to parse as JSON with tool_calls key
    try:
        data = json.loads(cleaned)
        if isinstance(data, dict) and "tool_calls" in data:
            calls = data["tool_calls"]
            reply_text = data.get("reply", "")
            return (calls if isinstance(calls, list) else [], reply_text)
    except (json.JSONDecodeError, ValueError):
        pass

    # Try to find embedded JSON object in the response
    brace_start = response.find('{"tool_calls"')
    if brace_start >= 0:
        # Find matching closing brace
        depth = 0
        for i in range(brace_start, len(response)):
            if response[i] == "{":
                depth += 1
            elif response[i] == "}":
                depth -= 1
                if depth == 0:
                    try:
                        data = json.loads(response[brace_start : i + 1])
                        calls = data.get("tool_calls", [])
                        reply_text = data.get("reply", response[:brace_start].strip())
                        return (calls if isinstance(calls, list) else [], reply_text)
                    except json.JSONDecodeError:
                        break

    return ([], "")


@router.post("/{proposal_id}/chat")
async def chat_with_proposal(
    proposal_id: str,
    body: ProposalChatRequest,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Chat with Quinn to refine a proposal using MCP tools.

    Quinn can use registered MCP tools (update_section, add_section,
    remove_section, list_sections, send_email, etc.) to make targeted
    edits to the proposal and interact with Gmail.
    """
    from app.services.llm import call_llm
    from app.mcp.registry import full_registry, ensure_setup
    ensure_setup()

    proposal = db.query(Proposal).filter(Proposal.id == proposal_id).first()
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    # Build conversation history
    history_text = ""
    for msg in body.history:
        role_label = "USER" if msg.role == "user" else "QUINN"
        history_text += f"{role_label}: {msg.content}\n"

    # Get MCP tool descriptions for the prompt
    tools_prompt = full_registry.get_tools_for_prompt()

    prompt = f"""You are Quinn, the AI assistant for DealMind. You help users refine their business proposals through conversation.

CURRENT PROPOSAL ID: {proposal_id}

CURRENT PROPOSAL CONTENT:
---
{proposal.content}
---

CONVERSATION HISTORY:
{history_text}

USER REQUEST: {body.message}

AVAILABLE MCP TOOLS:
{tools_prompt}

INSTRUCTIONS:
You have two ways to respond:

1. **Use tools** — If the user's request involves editing the proposal or interacting with email, respond with JSON:
```json
{{
  "tool_calls": [
    {{"name": "tool_name", "arguments": {{...}}}}
  ],
  "reply": "Brief explanation of what you're doing"
}}
```
IMPORTANT: For tool arguments, the proposal_id is "{proposal_id}". Do NOT include proposal_id in the arguments — it is automatically injected.

2. **Direct response** — If the user is asking a question or the request doesn't need tools, respond with:
```json
{{
  "reply": "Your conversational response here"
}}
```

RULES:
- Always respond with valid JSON (no extra text outside the JSON)
- For update_section: new_content should be the section BODY only (no heading line)
- For add_section: provide section_name, content, and optionally after_section
- You can call multiple tools at once
- Do not wrap JSON in code blocks"""

    tools_used = []

    try:
        raw = call_llm(prompt, max_tokens=8192)
        tool_calls, reply_text = _extract_tool_calls(raw)

        if tool_calls:
            # ── Execute MCP tool calls ──
            tool_results = []
            for tc in tool_calls:
                name = tc.get("name", "")
                args = tc.get("arguments", {})

                # Auto-inject proposal_id if tool needs it
                if "proposal_id" not in args:
                    args["proposal_id"] = proposal_id

                result = await full_registry.execute(
                    name,
                    args,
                    context={
                        "proposal_id": proposal_id,
                        "user_id": current_user.id,
                        "db": db,
                    },
                )
                tool_results.append({"tool": name, **result})
                tools_used.append(name)

            # If we don't have a reply from the initial response, ask LLM to summarize
            if not reply_text:
                summary_prompt = f"""You just executed these MCP tools on a proposal:

Tool Results:
{json.dumps(tool_results, indent=2, default=str)}

User's original request was: "{body.message}"

Write a brief, friendly reply (1-2 sentences) explaining what you did. Respond with plain text only, no JSON."""

                reply_text = call_llm(summary_prompt, max_tokens=512)

            # Check if any tool had errors
            errors = [r for r in tool_results if r.get("status") == "error"]
            if errors:
                error_msgs = [r.get("error", "Unknown error") for r in errors]
                reply_text = f"{reply_text}\n\n⚠️ Some operations had issues: {'; '.join(error_msgs)}"

        else:
            # No tool calls — use the LLM response directly
            # Try to extract reply from JSON format
            cleaned = _clean_json_block(raw)
            try:
                data = json.loads(cleaned)
                if isinstance(data, dict):
                    reply_text = data.get("reply", raw.strip())
                    # Check for updated_content (legacy fallback)
                    if "updated_content" in data:
                        proposal.content = data["updated_content"]
                        db.commit()
                else:
                    reply_text = raw.strip()
            except json.JSONDecodeError:
                reply_text = raw.strip()

        # Refresh proposal from DB (tools may have modified it)
        db.refresh(proposal)

        return {
            "reply": reply_text or "Done!",
            "updated_content": proposal.content,
            "tools_used": tools_used,
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Chat failed: {str(e)}")


@router.get("/{proposal_id}/export/docx")
def export_proposal_docx(
    proposal_id: str,
    template_id: str = Query("modern"),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Export proposal as a professionally formatted DOCX file with selectable template."""
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    # ── Validate template ──
    if template_id not in TEMPLATES:
        raise HTTPException(status_code=400, detail=f"Unknown template '{template_id}'. Available: {', '.join(TEMPLATES.keys())}")
    tc = TEMPLATES[template_id]

    proposal = db.query(Proposal).filter(Proposal.id == proposal_id).first()
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    deal = db.query(Deal).filter(Deal.id == proposal.deal_id).first()

    # Load team assignments and requirements for the deal
    from app.models.assignment import DealAssignment
    from app.models.deal import DealRequirement, DealAnalysis
    from app.models.employee import Employee
    assignments = []
    requirements = []
    analysis = None
    if deal:
        raw_assignments = (
            db.query(DealAssignment, Employee)
            .join(Employee, DealAssignment.employee_id == Employee.id)
            .filter(DealAssignment.deal_id == deal.id)
            .all()
        )
        assignments = []
        for a, emp in raw_assignments:
            a.employee_name = emp.name
            a.employee_role = emp.role or emp.department or "Team Member"
            assignments.append(a)
        requirements = db.query(DealRequirement).filter(DealRequirement.deal_id == deal.id).all()
        analysis = db.query(DealAnalysis).filter(DealAnalysis.deal_id == deal.id).order_by(DealAnalysis.created_at.desc()).first()

    doc = DocxDocument()

    # ═══════════════════════════════════════════════════
    # TEMPLATE COLORS (from config)
    # ═══════════════════════════════════════════════════
    PRIMARY = RGBColor(*tc["primary_rgb"])
    PRIMARY_LIGHT = RGBColor(*tc["primary_light_rgb"])
    ACCENT = RGBColor(*tc["accent_rgb"])
    DARK = RGBColor(*tc["dark_rgb"])
    GREY = RGBColor(*tc["grey_rgb"])
    LIGHT_GREY = RGBColor(*tc["light_grey_rgb"])
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    HDR_TEXT = RGBColor(*tc["table_header_text_rgb"])

    h_font = tc["heading_font"]
    b_font = tc["body_font"]

    client_name = deal.client_name if deal else "Valued Client"
    project_title = deal.title if deal else proposal.title or "Technical Proposal"

    # ═══════════════════════════════════════════════════
    # PAGE SETUP
    # ═══════════════════════════════════════════════════
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ═══════════════════════════════════════════════════
    # STYLES
    # ═══════════════════════════════════════════════════
    style = doc.styles["Normal"]
    style.font.name = b_font
    style.font.size = Pt(11)
    style.font.color.rgb = GREY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.25

    for level, (size, color) in {1: (22, PRIMARY), 2: (16, DARK), 3: (13, PRIMARY_LIGHT)}.items():
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = h_font
        h_style.font.size = Pt(size)
        h_style.font.bold = True
        h_style.font.color.rgb = color
        if level == 1:
            h_style.paragraph_format.space_before = Pt(30)
            h_style.paragraph_format.space_after = Pt(12)
        elif level == 2:
            h_style.paragraph_format.space_before = Pt(24)
            h_style.paragraph_format.space_after = Pt(10)
        else:
            h_style.paragraph_format.space_before = Pt(18)
            h_style.paragraph_format.space_after = Pt(8)

    # ═══════════════════════════════════════════════════
    # HELPER: Add a colored horizontal rule
    # ═══════════════════════════════════════════════════
    def add_colored_rule(color_hex=None, thickness=None):
        color_hex = color_hex or tc["primary_hex"]
        thickness = thickness or tc["rule_thickness"]
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(12)
        pPr = p._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="{thickness * 4}" w:space="1" w:color="{color_hex}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
        return p

    # ═══════════════════════════════════════════════════
    # HELPER: Styled info table (key-value pairs)
    # ═══════════════════════════════════════════════════
    def add_info_table(data_pairs):
        """Create a borderless info table with template-colored labels."""
        table = doc.add_table(rows=len(data_pairs), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[0].width = Inches(2)
        table.columns[1].width = Inches(4.5)
        for i, (label, value) in enumerate(data_pairs):
            cell_l = table.rows[i].cells[0]
            cell_l.text = ""
            p = cell_l.paragraphs[0]
            run = p.add_run(label)
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = PRIMARY
            run.font.name = b_font
            cell_r = table.rows[i].cells[1]
            cell_r.text = ""
            p = cell_r.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(11)
            run.font.color.rgb = DARK
            run.font.name = b_font
            for cell in [cell_l, cell_r]:
                tcEl = cell._tc
                tcPr = tcEl.get_or_add_tcPr()
                tcBorders = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                    f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                    f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                    f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                    f'</w:tcBorders>'
                )
                tcPr.append(tcBorders)
        doc.add_paragraph()

    # ═══════════════════════════════════════════════════
    # HELPER: Styled data table with header
    # ═══════════════════════════════════════════════════
    def add_styled_table(headers, rows_data, col_widths=None):
        table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_row = table.rows[0]
        for i, header_text in enumerate(headers):
            cell = hdr_row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(header_text)
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = HDR_TEXT
            run.font.name = b_font
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{tc["table_header_bg"]}" w:val="clear"/>')
            cell._tc.get_or_add_tcPr().append(shading)
            _set_cell_margins(cell)

        for r_idx, row_data in enumerate(rows_data):
            row = table.rows[r_idx + 1]
            bg = tc["table_row_alt_bg"] if r_idx % 2 == 0 else "FFFFFF"
            for c_idx, value in enumerate(row_data):
                cell = row.cells[c_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(str(value))
                run.font.size = Pt(10)
                run.font.color.rgb = GREY
                run.font.name = b_font
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shading)
                _set_cell_margins(cell)

        if col_widths:
            for i, w in enumerate(col_widths):
                table.columns[i].width = Inches(w)

        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        bdr_hex = tc["table_border_hex"]
        bdr_in = tc["table_border_inside"]
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{bdr_hex}"/>'
            f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{bdr_hex}"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{bdr_hex}"/>'
            f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{bdr_hex}"/>'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{bdr_in}"/>'
            f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{bdr_in}"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

        doc.add_paragraph()
        return table

    def _set_cell_margins(cell):
        tcEl = cell._tc
        tcPr = tcEl.get_or_add_tcPr()
        tcMar = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'  <w:top w:w="60" w:type="dxa"/>'
            f'  <w:left w:w="100" w:type="dxa"/>'
            f'  <w:bottom w:w="60" w:type="dxa"/>'
            f'  <w:right w:w="100" w:type="dxa"/>'
            f'</w:tcMar>'
        )
        tcPr.append(tcMar)

    # ═══════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════

    # Top accent bar
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{tc["cover_border_sz"]}" w:space="0" w:color="{tc["primary_hex"]}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # Company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(80)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("ESSHVA")
    run.font.size = Pt(42)
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    run.font.name = h_font

    # Tagline
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("Technology Solutions & Innovation")
    run.font.size = Pt(14)
    run.font.color.rgb = LIGHT_GREY
    run.font.name = b_font
    run.font.italic = True

    add_colored_rule(tc["primary_hex"], 3)

    # Proposal title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(16)
    run = p.add_run(proposal.title or "Technical Proposal")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = DARK
    run.font.name = h_font

    # Subtitle with project name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(f"For {project_title}")
    run.font.size = Pt(16)
    run.font.color.rgb = PRIMARY_LIGHT
    run.font.name = b_font

    # Cover page info table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)

    cover_data = [
        ("CLIENT", client_name),
        ("PROJECT", project_title),
        ("DATE", datetime.utcnow().strftime("%B %d, %Y")),
        ("REFERENCE", f"ESSHVA-{(proposal.id or '')[:8].upper()}"),
        ("VERSION", f"v{proposal.version}"),
    ]

    add_info_table(cover_data)

    # Contact info at bottom
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)
    run = p.add_run("ESSHVA  ")
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    run.font.name = b_font
    run = p.add_run("|  contact@esshva.com  |  www.esshva.com")
    run.font.size = Pt(10)
    run.font.color.rgb = LIGHT_GREY
    run.font.name = b_font

    # Confidential notice
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("CONFIDENTIAL — This document is the property of ESSHVA and is intended solely for the named recipient.")
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = LIGHT_GREY
    run.font.name = b_font

    doc.add_page_break()

    # ═══════════════════════════════════════════════════
    # HEADERS & FOOTERS
    # ═══════════════════════════════════════════════════
    section = doc.sections[-1]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("ESSHVA")
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    run.font.name = b_font
    run = hp.add_run(f"  |  {proposal.title or 'Proposal'}  |  {client_name}")
    run.font.size = Pt(8)
    run.font.color.rgb = LIGHT_GREY
    run.font.name = b_font

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("CONFIDENTIAL  |  ")
    run.font.size = Pt(8)
    run.font.color.rgb = LIGHT_GREY
    run.font.name = b_font
    run = fp.add_run("ESSHVA")
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    run.font.name = b_font
    run = fp.add_run(f"  |  {datetime.utcnow().strftime('%B %Y')}")
    run.font.size = Pt(8)
    run.font.color.rgb = LIGHT_GREY
    run.font.name = b_font

    # ═══════════════════════════════════════════════════
    # TABLE OF CONTENTS — Quick Reference
    # ═══════════════════════════════════════════════════
    doc.add_heading("Project Overview", level=1)
    add_colored_rule()

    overview_data = [
        ("Client", client_name),
        ("Project", project_title),
    ]
    if deal and deal.description:
        overview_data.append(("Description", deal.description[:200]))
    if proposal.compliance_score is not None:
        score_pct = f"{proposal.compliance_score * 100:.0f}%"
        overview_data.append(("Compliance Score", score_pct))
    if assignments:
        overview_data.append(("Team Size", f"{len(assignments)} members assigned"))

    add_info_table(overview_data)

    # ═══════════════════════════════════════════════════
    # MAIN PROPOSAL CONTENT
    # ═══════════════════════════════════════════════════
    content = proposal.content or ""
    lines = content.split("\n")
    dark_rgb_tuple = tc["dark_rgb"]

    def strip_bold_wrap(text):
        """Remove wrapping ** from text like '**Executive Summary**' → 'Executive Summary'."""
        t = text.strip()
        if t.startswith("**") and t.endswith("**") and t.count("**") == 2:
            t = t[2:-2].strip()
        return t

    def is_table_separator(text):
        """Check if a line is a markdown table separator like |:---|:---|."""
        return bool(re.match(r"^\|[\s:_-]+(\|[\s:_-]+)+\|?\s*$", text))

    def parse_md_table(table_lines):
        """Parse markdown table lines into header + rows for add_styled_table."""
        parsed_rows = []
        for tl in table_lines:
            tl = tl.strip()
            if is_table_separator(tl):
                continue
            cells = [c.strip() for c in tl.strip("|").split("|")]
            cells = [strip_bold_wrap(c) for c in cells]
            parsed_rows.append(cells)
        if not parsed_rows:
            return [], []
        headers = parsed_rows[0]
        data = parsed_rows[1:]
        return headers, data

    # Pre-process: collect lines and detect table blocks
    i = 0
    blank_count = 0
    h1_count = 0  # Track H1s to add page breaks before major sections

    while i < len(lines):
        stripped = lines[i].strip()

        # ── Blank lines → controlled spacing ──
        if not stripped:
            blank_count += 1
            # Only add one spacer per blank-line group (prevents excessive gaps)
            if blank_count == 1:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        blank_count = 0

        # ── Markdown table block ──
        if "|" in stripped and stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            headers, data = parse_md_table(table_lines)
            if headers and data:
                # Calculate column widths to fit page
                n_cols = len(headers)
                avail = 6.5  # inches of content width
                col_w = [round(avail / n_cols, 1)] * n_cols
                add_styled_table(headers, data, col_widths=col_w)
            else:
                # Fallback: render as regular text
                for tl in table_lines:
                    p = doc.add_paragraph()
                    _add_formatted_runs(p, tl.strip(), dark_rgb_tuple)
            continue

        # ── Table separator line on its own (skip) ──
        if is_table_separator(stripped):
            i += 1
            continue

        # ── Heading 1 ──
        if stripped.startswith("# ") and not stripped.startswith("## "):
            heading_text = strip_bold_wrap(stripped[2:].strip())
            # Also strip numbering-style prefix from headings
            heading_text = re.sub(r"^\d+\.\s*", "", heading_text).strip()
            h1_count += 1
            # Add page break before major sections (skip the very first H1 — it follows the overview)
            if h1_count > 1:
                doc.add_page_break()
            doc.add_heading(heading_text, level=1)
            add_colored_rule()
        # ── Heading 2 ──
        elif stripped.startswith("## ") and not stripped.startswith("### "):
            heading_text = strip_bold_wrap(stripped[3:].strip())
            heading_text = re.sub(r"^\d+\.\s*", "", heading_text).strip()
            doc.add_heading(heading_text, level=2)
        # ── Heading 3 ──
        elif stripped.startswith("### "):
            heading_text = strip_bold_wrap(stripped[4:].strip())
            heading_text = re.sub(r"^\d+\.\s*", "", heading_text).strip()
            doc.add_heading(heading_text, level=3)
        # ── Horizontal rule ──
        elif stripped in ("---", "***", "___"):
            add_colored_rule(tc["table_border_hex"], 1)
        # ── Bullet list ──
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            _add_formatted_runs(p, text, dark_rgb_tuple)
        # ── Numbered list ──
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            _add_formatted_runs(p, text, dark_rgb_tuple)
        # ── Regular paragraph ──
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            _add_formatted_runs(p, stripped, dark_rgb_tuple)

        i += 1

    # ═══════════════════════════════════════════════════
    # NOTE: Team members and compliance info are already included
    # in the LLM-generated proposal content above. No separate
    # appendix sections are needed — avoids duplicate content.
    # ═══════════════════════════════════════════════════

    # ═══════════════════════════════════════════════════
    # CLOSING PAGE
    # ═══════════════════════════════════════════════════
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(100)
    run = p.add_run("ESSHVA")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    run.font.name = h_font

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Thank you for the opportunity to partner with you.")
    run.font.size = Pt(14)
    run.font.color.rgb = GREY
    run.font.name = b_font

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(f"We look forward to delivering exceptional results for {client_name}.")
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = LIGHT_GREY
    run.font.name = b_font

    add_colored_rule(tc["primary_hex"], 2)

    # Contact block
    contact_items = [
        ("Email", "contact@esshva.com"),
        ("Website", "www.esshva.com"),
    ]
    for label, value in contact_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}: ")
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = PRIMARY
        run.font.name = b_font
        run = p.add_run(value)
        run.font.size = Pt(11)
        run.font.color.rgb = GREY
        run.font.name = b_font

    # Final confidential notice
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    run = p.add_run("CONFIDENTIAL")
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = LIGHT_GREY
    run.font.name = b_font

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"This document is the property of ESSHVA and is prepared exclusively for {client_name}.")
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = LIGHT_GREY
    run.font.name = b_font

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Reference: ESSHVA-{(proposal.id or '')[:8].upper()}  |  {datetime.utcnow().strftime('%B %d, %Y')}")
    run.font.size = Pt(8)
    run.font.color.rgb = LIGHT_GREY
    run.font.name = b_font

    # ═══════════════════════════════════════════════════
    # SAVE
    # ═══════════════════════════════════════════════════
    os.makedirs("exports", exist_ok=True)
    safe_title = re.sub(r"[^\w\s-]", "", proposal.title or "proposal").strip().replace(" ", "_")[:50]
    safe_client = re.sub(r"[^\w\s-]", "", client_name).strip().replace(" ", "_")[:30]
    filename = f"ESSHVA_Proposal_{safe_client}_{safe_title}.docx"
    filepath = os.path.join("exports", filename)
    doc.save(filepath)

    return FileResponse(
        path=filepath,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _add_formatted_runs(paragraph, text: str, dark_rgb_tuple=(0x1a, 0x1a, 0x2e)):
    """Parse markdown bold/italic inline and add as runs to a paragraph."""
    from docx.shared import RGBColor
    # Match: ***bold+italic***, **bold**, *italic*
    parts = re.split(r"(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
            run.font.color.rgb = RGBColor(*dark_rgb_tuple)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.color.rgb = RGBColor(*dark_rgb_tuple)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


@router.post("/{proposal_id}/export")
def export_proposal(
    proposal_id: str,
    export_request: ProposalExportRequest,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user),
):
    """Export proposal to file (legacy endpoint)."""
    proposal = db.query(Proposal).filter(Proposal.id == proposal_id).first()
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")
    return {"status": "processing", "message": f"Use GET /export/docx for Word export"}


PROPOSAL_UPLOAD_DIR = "uploads/proposals"
os.makedirs(PROPOSAL_UPLOAD_DIR, exist_ok=True)

ALLOWED_EXTENSIONS = {"pdf", "docx", "xlsx", "txt", "pptx", "csv", "png", "jpg", "jpeg"}


@router.post("/{proposal_id}/attachments")
async def upload_proposal_attachment(
    proposal_id: str,
    file: UploadFile = File(...),
    background_tasks: BackgroundTasks = BackgroundTasks(),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Upload a supporting document/attachment to a proposal. Auto-indexes text files into RAG."""
    proposal = db.query(Proposal).filter(Proposal.id == proposal_id).first()
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    ext = file.filename.rsplit(".", 1)[1].lower() if "." in file.filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"File type .{ext} not allowed")

    # Save file
    proposal_dir = os.path.join(PROPOSAL_UPLOAD_DIR, proposal_id)
    os.makedirs(proposal_dir, exist_ok=True)

    doc_id = str(uuid4())
    file_path = os.path.join(proposal_dir, f"{doc_id}_{file.filename}")

    async with aiofiles.open(file_path, "wb") as f:
        content = await file.read()
        await f.write(content)

    # Create document record linked to the proposal's deal
    document = Document(
        id=doc_id,
        deal_id=proposal.deal_id,
        filename=file.filename,
        original_filename=file.filename,
        file_path=file_path,
        file_type=ext,
        file_size=len(content),
        doc_category="proposal",
        is_processed=False,
        uploaded_by=current_user.id,
    )
    db.add(document)
    db.commit()
    db.refresh(document)

    # Process in background (extract text + index into RAG "proposals" collection)
    from app.routers.documents import process_document_background
    background_tasks.add_task(process_document_background, doc_id, file_path, "proposal", db)

    return {
        "document_id": doc_id,
        "filename": file.filename,
        "file_size": len(content),
        "message": "File uploaded and queued for processing",
    }


@router.get("/{proposal_id}/attachments")
def list_proposal_attachments(
    proposal_id: str,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """List all files attached to a proposal (doc_category='proposal' for the deal)."""
    proposal = db.query(Proposal).filter(Proposal.id == proposal_id).first()
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    documents = (
        db.query(Document)
        .filter(Document.deal_id == proposal.deal_id, Document.doc_category == "proposal")
        .order_by(Document.created_at.desc())
        .all()
    )

    return [
        {
            "id": doc.id,
            "filename": doc.original_filename or doc.filename,
            "file_type": doc.file_type,
            "file_size": doc.file_size,
            "is_processed": doc.is_processed,
            "created_at": doc.created_at.isoformat() if doc.created_at else None,
        }
        for doc in documents
    ]


@router.delete("/{proposal_id}/attachments/{document_id}")
def delete_proposal_attachment(
    proposal_id: str,
    document_id: str,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """Delete an attachment from a proposal."""
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    if os.path.exists(document.file_path):
        os.remove(document.file_path)

    db.delete(document)
    db.commit()

    return {"message": "Attachment deleted"}
