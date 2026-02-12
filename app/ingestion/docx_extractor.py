from docx import Document as DocxDocument
from pathlib import Path


class DocxExtractor:
    """Extract text and structure from Word documents."""

    def extract(self, file_path: str) -> dict:
        doc = DocxDocument(file_path)
        result = {
            "text": "",
            "sections": [],
            "tables_found": 0,
            "page_count": None,  # DOCX doesn't have reliable page count
            "metadata": {},
        }

        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                result["sections"].append({
                    "title": para.text.strip(),
                    "style": para.style.name,
                })

        result["text"] = "\n".join(text_parts)

        # Extract tables
        for table in doc.tables:
            result["tables_found"] += 1

        # Core properties
        if doc.core_properties:
            result["metadata"] = {
                "author": doc.core_properties.author or "",
                "title": doc.core_properties.title or "",
                "created": str(doc.core_properties.created) if doc.core_properties.created else "",
            }

        return result


docx_extractor = DocxExtractor()
